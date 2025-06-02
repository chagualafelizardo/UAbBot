import os
from docx import Document
from typing import List, Dict

class CursoExtractor:
    def __init__(self):
        self.file_path = r'C:\Users\Felizardo Chaguala\Desktop\Rasa\UAbBot\rasa.uabbot.files\conhecer-a-uab\conhecer-a-uab.docx'
    
    def extract_cursos_por_tipo(self) -> Dict[str, List[str]]:
        try:
            doc = Document(self.file_path)
            cursos_por_tipo = {}
            tipo_atual = None
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Verifica se é título de seção
                if text.lower().startswith('cursos de'):
                    tipo_atual = text.replace("Cursos de", "").strip().lower()
                    cursos_por_tipo[tipo_atual] = []
                    continue

                # Adiciona cursos ao grupo atual
                if tipo_atual and self.is_complete_curso(text):
                    cursos_por_tipo[tipo_atual].append(text.lstrip("• ").strip())

            # Também verifica se há cursos em tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if not cell_text:
                            continue

                        if cell_text.lower().startswith("cursos de"):
                            tipo_atual = cell_text.replace("Cursos de", "").strip().lower()
                            cursos_por_tipo[tipo_atual] = []
                        elif tipo_atual and self.is_complete_curso(cell_text):
                            cursos_por_tipo[tipo_atual].append(cell_text.lstrip("• ").strip())

            return cursos_por_tipo

        except Exception as e:
            print(f"\nERRO: Falha ao ler o arquivo: {str(e)}")
            return {}

    def is_complete_curso(self, text: str) -> bool:
        tipos = ['licenciatura', 'mestrado', 'doutoramento']
        return any(tipo in text.lower() for tipo in tipos)

def main():
    print("=== Extrator de Cursos da UAb por Tipo ===")
    extractor = CursoExtractor()
    cursos_por_tipo = extractor.extract_cursos_por_tipo()

    if not cursos_por_tipo:
        print("Nenhum curso encontrado.")
        return

    tipos_disponiveis = list(cursos_por_tipo.keys())
    print("\nTipos de curso disponíveis:")
    for tipo in tipos_disponiveis:
        print(f"- {tipo.capitalize()}")

    while True:
        escolha = input("\nDigite o nome do tipo de curso (ou 'sair' para terminar): ").strip().lower()
        if escolha == 'sair':
            print("Encerrando.")
            break
        elif escolha in cursos_por_tipo:
            cursos = cursos_por_tipo[escolha]
            print(f"\nForam encontrados {len(cursos)} cursos de {escolha.capitalize()}:")
            for i, curso in enumerate(cursos, 1):
                print(f"{i}. {curso}")
        else:
            print("Tipo inválido. Tente novamente.")

if __name__ == "__main__":
    main()
